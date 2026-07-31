import base64
import json
import os
import sys

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


def load_params(raw):
    padding = "=" * (-len(raw) % 4)
    return json.loads(base64.urlsafe_b64decode(raw + padding).decode("utf-8"))


def all_paragraphs(doc):
    values = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                values.extend(cell.paragraphs)
    return values


def replace_in_runs(paragraph, old, new):
    full = "".join(run.text for run in paragraph.runs)
    start = full.find(old)
    if start < 0:
        return False
    end = start + len(old)
    cursor = 0
    start_run = end_run = 0
    start_offset = end_offset = 0
    for index, run in enumerate(paragraph.runs):
        next_cursor = cursor + len(run.text)
        if cursor <= start < next_cursor:
            start_run, start_offset = index, start - cursor
        if cursor < end <= next_cursor:
            end_run, end_offset = index, end - cursor
            break
        cursor = next_cursor
    prefix = paragraph.runs[start_run].text[:start_offset]
    suffix = paragraph.runs[end_run].text[end_offset:]
    paragraph.runs[start_run].text = prefix + new
    for index in range(start_run + 1, end_run):
        paragraph.runs[index].text = ""
    if end_run != start_run:
        paragraph.runs[end_run].text = suffix
    else:
        paragraph.runs[start_run].text += suffix
    return True


def insert_after(paragraph, text):
    node = OxmlElement("w:p")
    paragraph._p.addnext(node)
    Paragraph(node, paragraph._parent).add_run(text)


def process(input_path, output_path, params):
    doc = Document(input_path)
    paragraphs = all_paragraphs(doc)
    action = params.get("action", "replace_text")
    scope = params.get("scope", "first")
    target = params.get("target_text", "")
    replacement = "" if action == "delete_text" else params.get("new_text", "")
    changed = 0

    if action == "insert_text":
        text = params.get("text", "")
        index = params.get("paragraph_index")
        anchor = params.get("anchor_text", "")
        selected = paragraphs[index] if isinstance(index, int) and 0 <= index < len(paragraphs) else None
        if selected is None and anchor:
            selected = next((p for p in paragraphs if anchor in p.text), None)
        if selected is None:
            doc.add_paragraph(text)
        else:
            insert_after(selected, text)
        changed = 1
    else:
        if not target:
            raise RuntimeError("target_text is required")
        for paragraph in paragraphs:
            while target in paragraph.text:
                if not replace_in_runs(paragraph, target, replacement):
                    break
                changed += 1
                if scope != "all":
                    break
            if changed and scope != "all":
                break

    if changed == 0:
        raise RuntimeError("No matching edit position found")
    doc.save(output_path)


if __name__ == "__main__":
    try:
        if len(sys.argv) < 4 or not os.path.isfile(sys.argv[1]):
            raise RuntimeError("Invalid input or parameters")
        process(sys.argv[1], sys.argv[2], load_params(sys.argv[3]))
        if not os.path.isfile(sys.argv[2]) or os.path.getsize(sys.argv[2]) == 0:
            raise RuntimeError("Output file generation failed")
    except Exception as error:
        print(f"[Execution Error]: {error}", file=sys.stderr)
        sys.exit(1)
