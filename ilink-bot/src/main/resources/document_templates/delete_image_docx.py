import sys
import os
from docx import Document

if len(sys.argv) < 3:
    print("[Error] Usage: script.py input.docx output.docx [page_or_all]", file=sys.stderr)
    sys.exit(1)

input_path = sys.argv[1]
output_path = sys.argv[2]
target = sys.argv[3] if len(sys.argv) > 3 else "all"  # all|页码


def paragraph_pages(paragraphs):
    result = {}
    page = 1
    reliable = False
    for index, paragraph in enumerate(paragraphs):
        if index > 0 and paragraph._p.xpath(
                './/w:pPr/w:pageBreakBefore[not(@w:val="0") and not(@w:val="false")]'):
            page += 1
            reliable = True
        result[id(paragraph._element)] = page
        breaks = paragraph._p.xpath('.//w:lastRenderedPageBreak | .//w:br[@w:type="page"]')
        if breaks:
            page += len(breaks)
            reliable = True
    return result, reliable

def delete_images(inp, outp, target_scope):
    doc = Document(inp)
    removed = 0
    pages, reliable_pages = paragraph_pages(doc.paragraphs)
    target_page = int(target_scope) if target_scope.isdigit() else None
    if target_page is not None and target_page > 1 and not reliable_pages:
        raise RuntimeError("DOCX 未保存可用分页边界，无法按页码删除图片")
    
    # 删除段落中的图片
    for para in doc.paragraphs:
        paragraph_page = pages.get(id(para._element), 1)
        for run in para.runs:
            drawings = run._r.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
            if not drawings:
                drawings = run._r.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
            inline_shapes = run._r.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict')
            all_imgs = drawings + inline_shapes
            if all_imgs:
                if target_scope == "all" or target_page == paragraph_page:
                    run._r.getparent().remove(run._r)
                    removed += 1
    
    # 删除表格中的图片
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        drawings = run._r.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                        inline_shapes = run._r.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict')
                        if drawings or inline_shapes:
                            if target_scope == "all":
                                run._r.getparent().remove(run._r)
                                removed += 1

    if removed == 0:
        if target_page is not None:
            raise RuntimeError(f"第 {target_page} 页没有找到可删除的图片")
        raise RuntimeError("文档中没有找到可删除的图片")
    
    doc.save(outp)

if __name__ == "__main__":
    try:
        if not os.path.exists(input_path):
            raise FileNotFoundError(f"Input file not found: {input_path}")
        delete_images(input_path, output_path, target)
        if not os.path.exists(output_path) or os.path.getsize(output_path) == 0:
            raise RuntimeError("Output file generation failed")
        sys.exit(0)
    except Exception as e:
        import traceback
        print(f"[Execution Error]: {e}", file=sys.stderr)
        traceback.print_exc(file=sys.stderr)
        sys.exit(1)
