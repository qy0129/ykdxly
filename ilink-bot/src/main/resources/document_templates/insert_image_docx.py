import ast
import base64
import json
import os
import sys

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches


if len(sys.argv) < 4:
    print("[Error] Usage: script.py input.docx output.docx image.png [params_json]", file=sys.stderr)
    sys.exit(1)

input_path = sys.argv[1]
output_path = sys.argv[2]
image_path = sys.argv[3]

def load_params(raw):
    if not raw:
        return {"position": "end"}
    if raw.lstrip().startswith("{"):
        try:
            return json.loads(raw)
        except json.JSONDecodeError:
            try:
                value = ast.literal_eval(raw)
                if isinstance(value, dict):
                    return value
            except (SyntaxError, ValueError):
                pass
            if "=" in raw:
                result = {}
                for item in raw.strip()[1:-1].split(","):
                    key, value = item.split("=", 1)
                    value = value.strip()
                    result[key.strip()] = int(value) if value.isdigit() else value
                return result
            raise
    padded = raw + "=" * (-len(raw) % 4)
    decoded = base64.urlsafe_b64decode(padded).decode("utf-8")
    return json.loads(decoded)


params = load_params(sys.argv[4]) if len(sys.argv) > 4 else {"position": "end"}


def has_image(paragraph):
    return any(run._r.xpath(".//w:drawing") or run._r.xpath(".//w:pict") for run in paragraph.runs)


def image_paragraphs(document):
    return [paragraph for paragraph in document.paragraphs if has_image(paragraph)]


def new_image_paragraph(document, image, values=None):
    paragraph = document.add_paragraph()
    width_inches = 4
    if values:
        if values.get("width_inches") is not None:
            width_inches = float(values["width_inches"])
        elif values.get("width_cm") is not None:
            width_inches = float(values["width_cm"]) / 2.54
    paragraph.add_run().add_picture(image, width=Inches(width_inches))
    return paragraph


def append_with_warning(document, image, warning, values=None):
    new_image_paragraph(document, image, values)
    print(f"[Warning] {warning}，图片已追加到文档末尾。")


def page_break_before(paragraph):
    return bool(paragraph._p.xpath(
        './/w:pPr/w:pageBreakBefore[not(@w:val="0") and not(@w:val="false")]'
    ))


def page_break_count(paragraph):
    return len(paragraph._p.xpath(
        './/w:lastRenderedPageBreak | .//w:br[@w:type="page"]'
    ))


def paragraphs_by_page(paragraphs):
    pages = {}
    page = 1
    reliable = False
    for index, paragraph in enumerate(paragraphs):
        if index > 0 and page_break_before(paragraph):
            page += 1
            reliable = True
        pages.setdefault(page, []).append(paragraph)
        breaks = page_break_count(paragraph)
        if breaks:
            page += breaks
            reliable = True
    return pages, reliable


def paragraph_with_text(paragraphs, anchor, occurrence, page_paragraphs=None):
    candidates = page_paragraphs if page_paragraphs is not None else paragraphs
    remaining = max(1, int(occurrence or 1))
    for paragraph in candidates:
        count = paragraph.text.count(anchor)
        if remaining <= count:
            return paragraph
        remaining -= count
    return None


def insert_image(inp, outp, image, values):
    document = Document(inp)
    position = values.get("position", "end")
    paragraphs = document.paragraphs
    pages, reliable_pages = paragraphs_by_page(paragraphs)
    page = int(values.get("page", 1)) if values.get("page") is not None else None

    if position == "new_page":
        page_break = document.add_paragraph()
        page_break.add_run().add_break(WD_BREAK.PAGE)
        new_image_paragraph(document, image, values)
    elif position == "page_auto":
        target_page = page or 1
        page_paragraphs = pages.get(target_page)
        if page_paragraphs and (reliable_pages or target_page == 1):
            target = page_paragraphs[-1] if reliable_pages else page_paragraphs[0]
            inserted = new_image_paragraph(document, image, values)
            if reliable_pages and page_break_count(target) > 0:
                target._element.addprevious(inserted._element)
            else:
                target._element.addnext(inserted._element)
            if not reliable_pages:
                print("[Warning] DOCX 未保存分页边界，已将图片放在文档开头的第一页区域。")
        else:
            append_with_warning(document, image, f"无法可靠定位第 {target_page} 页", values)
    elif position in {"before_text", "after_text"}:
        anchor = values.get("anchor_text", "")
        occurrence = values.get("occurrence", 1)
        page_paragraphs = pages.get(page) if page is not None and reliable_pages else None
        target = paragraph_with_text(paragraphs, anchor, occurrence, page_paragraphs)
        if target is not None:
            inserted = new_image_paragraph(document, image, values)
            if position == "before_text":
                target._element.addprevious(inserted._element)
            else:
                target._element.addnext(inserted._element)
            if page is not None and not reliable_pages:
                print("[Warning] DOCX 未保存分页边界，本次按你引用的原文定位，未验证物理页码。")
        else:
            append_with_warning(document, image, "未找到你引用的原文位置", values)
    elif position in {"after_paragraph", "before_paragraph"}:
        index = int(values.get("paragraph_index", values.get("index", 0)))
        if 0 <= index < len(paragraphs):
            target = paragraphs[index]
            inserted = new_image_paragraph(document, image, values)
            if position == "before_paragraph":
                target._element.addprevious(inserted._element)
            else:
                target._element.addnext(inserted._element)
        else:
            append_with_warning(document, image, "未找到你指定的段落", values)
    elif position in {"before_image", "after_image"}:
        images = image_paragraphs(document)
        index = int(values.get("image_index", 1)) - 1
        if 0 <= index < len(images):
            target = images[index]
            inserted = new_image_paragraph(document, image, values)
            if position == "before_image":
                target._element.addprevious(inserted._element)
            else:
                target._element.addnext(inserted._element)
        else:
            append_with_warning(document, image, "未找到你指定的图片", values)
    elif position == "between_paragraph_image":
        paragraph_index = int(values.get("paragraph_index", 0))
        image_index = int(values.get("image_index", 1)) - 1
        images = image_paragraphs(document)
        if 0 <= paragraph_index < len(paragraphs) and 0 <= image_index < len(images):
            text_paragraph = paragraphs[paragraph_index]
            image_paragraph = images[image_index]
            body_children = list(document.element.body)
            text_order = body_children.index(text_paragraph._element)
            image_order = body_children.index(image_paragraph._element)
            if text_order < image_order:
                inserted = new_image_paragraph(document, image, values)
                image_paragraph._element.addprevious(inserted._element)
            else:
                append_with_warning(document, image, "指定文字不在目标图片之前", values)
        else:
            append_with_warning(document, image, "未找到你指定的文字或图片", values)
    else:
        new_image_paragraph(document, image, values)

    document.save(outp)


if __name__ == "__main__":
    try:
        if not os.path.exists(input_path):
            raise FileNotFoundError(f"Input file not found: {input_path}")
        if not os.path.exists(image_path):
            raise FileNotFoundError(f"Image file not found: {image_path}")
        insert_image(input_path, output_path, image_path, params)
        if not os.path.exists(output_path) or os.path.getsize(output_path) == 0:
            raise RuntimeError("Output file generation failed")
        sys.exit(0)
    except Exception as exc:
        import traceback
        print(f"[Execution Error]: {exc}", file=sys.stderr)
        traceback.print_exc(file=sys.stderr)
        sys.exit(1)
