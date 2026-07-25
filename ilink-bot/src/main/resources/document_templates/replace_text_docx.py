import sys
import os
from docx import Document

if len(sys.argv) < 5:
    print("[Error] Usage: script.py input.docx output.docx old_text new_text", file=sys.stderr)
    sys.exit(1)

input_path = sys.argv[1]
output_path = sys.argv[2]
old_text = sys.argv[3]
new_text = sys.argv[4]

def replace_in_runs(runs, old, new):
    replaced = 0
    while True:
        full_text = "".join(run.text for run in runs)
        start = full_text.find(old)
        if start < 0:
            return replaced

        end = start + len(old)
        cursor = 0
        start_run = start_offset = end_run = end_offset = None
        for index, run in enumerate(runs):
            next_cursor = cursor + len(run.text)
            if start_run is None and start < next_cursor:
                start_run = index
                start_offset = start - cursor
            if end_run is None and end <= next_cursor:
                end_run = index
                end_offset = end - cursor
                break
            cursor = next_cursor

        if start_run is None or end_run is None:
            return replaced
        if start_run == end_run:
            text = runs[start_run].text
            runs[start_run].text = text[:start_offset] + new + text[end_offset:]
        else:
            prefix = runs[start_run].text[:start_offset]
            suffix = runs[end_run].text[end_offset:]
            runs[start_run].text = prefix + new
            for index in range(start_run + 1, end_run):
                runs[index].text = ""
            runs[end_run].text = suffix
        replaced += 1


def replace_text(inp, outp, old, new):
    doc = Document(inp)
    replaced = 0
    
    # 替换段落中的文字
    for para in doc.paragraphs:
        if old in para.text:
            replaced += replace_in_runs(para.runs, old, new)
    
    # 替换表格中的文字
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if old in para.text:
                        replaced += replace_in_runs(para.runs, old, new)

    if replaced == 0:
        raise RuntimeError(f"Text not found: {old}")
    
    doc.save(outp)

if __name__ == "__main__":
    try:
        if not os.path.exists(input_path):
            raise FileNotFoundError(f"Input file not found: {input_path}")
        replace_text(input_path, output_path, old_text, new_text)
        if not os.path.exists(output_path) or os.path.getsize(output_path) == 0:
            raise RuntimeError("Output file generation failed")
        sys.exit(0)
    except Exception as e:
        import traceback
        print(f"[Execution Error]: {e}", file=sys.stderr)
        traceback.print_exc(file=sys.stderr)
        sys.exit(1)
